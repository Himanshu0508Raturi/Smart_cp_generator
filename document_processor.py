import os
import io
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from werkzeug.datastructures import FileStorage
import PyPDF2
from app import app

class DocumentProcessor:
    """Handles document processing operations"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def extract_text_from_file(self, file: FileStorage) -> str:
        """Extract text content from uploaded file"""
        try:
            filename = file.filename.lower()
            
            if filename.endswith('.txt'):
                return file.read().decode('utf-8')
            
            elif filename.endswith('.docx'):
                # Save file temporarily
                temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{datetime.now().timestamp()}.docx")
                file.save(temp_path)
                
                # Extract text from docx
                doc = Document(temp_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                # Clean up temp file
                os.remove(temp_path)
                return text
            
            elif filename.endswith('.pdf'):
                # Save file temporarily
                temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{datetime.now().timestamp()}.pdf")
                file.save(temp_path)
                
                # Extract text from PDF
                text = ""
                with open(temp_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                
                # Clean up temp file
                os.remove(temp_path)
                return text
            
            else:
                raise ValueError(f"Unsupported file type: {filename}")
                
        except Exception as e:
            self.logger.error(f"Error extracting text from file: {str(e)}")
            raise
    
    def merge_documents(self, document_contents: dict, extracted_clauses: dict) -> str:
        """Merge fixture recap, base CP, and negotiated clauses into final contract"""
        
        # Start with base CP if available, otherwise start fresh
        final_contract = document_contents.get('base_cp', '')
        
        # Add fixture recap information
        if document_contents.get('fixture_recap'):
            fixture_section = f"""
=== FIXTURE RECAP ===

{document_contents['fixture_recap']}

"""
            final_contract = fixture_section + final_contract
        
        # Process and integrate negotiated clauses
        if document_contents.get('negotiated_clauses'):
            negotiated_section = f"""

=== NEGOTIATED CLAUSES ===

{document_contents['negotiated_clauses']}

"""
            final_contract += negotiated_section
        
        # Add extracted clauses summary
        if extracted_clauses:
            clauses_summary = self._format_extracted_clauses(extracted_clauses)
            final_contract += f"""

=== EXTRACTED CLAUSES SUMMARY ===

{clauses_summary}

"""
        
        # Add contract footer
        final_contract += f"""

=== CONTRACT GENERATED ===
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
By: Smart CP Generator - Contract Automation System

This document combines the provided fixture recap, base charter party agreement, and negotiated clauses into a unified contract document. Please review all terms carefully before execution.
"""
        
        return final_contract
    
    def _format_extracted_clauses(self, extracted_clauses: dict) -> str:
        """Format extracted clauses for display"""
        formatted = ""
        
        for category, clauses in extracted_clauses.items():
            if clauses:
                formatted += f"\n{category.upper()}:\n"
                for i, clause in enumerate(clauses, 1):
                    formatted += f"{i}. {clause}\n"
                formatted += "\n"
        
        return formatted
    
    def generate_output_files(self, content: str, contract_id: int, contract_name: str) -> tuple:
        """Generate Word and PDF files from contract content"""
        
        # Generate safe filename
        safe_name = "".join(c for c in contract_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Generate DOCX file
        docx_filename = f"{safe_name}_{timestamp}.docx"
        docx_path = os.path.join(app.config['GENERATED_FOLDER'], docx_filename)
        self._create_docx(content, docx_path, contract_name)
        
        # Generate PDF file
        pdf_filename = f"{safe_name}_{timestamp}.pdf"
        pdf_path = os.path.join(app.config['GENERATED_FOLDER'], pdf_filename)
        self._create_pdf(content, pdf_path, contract_name)
        
        return docx_path, pdf_path
    
    def _create_docx(self, content: str, file_path: str, title: str):
        """Create Word document from content"""
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph("")  # Empty line
        
        # Split content into sections and paragraphs
        sections = content.split('===')
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
                
            # Check if this is a section header
            if section.isupper() and len(section) < 50:
                doc.add_heading(section, level=1)
            else:
                # Split into paragraphs
                paragraphs = section.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        # Check if it's a subsection header
                        if para.isupper() and len(para) < 100:
                            doc.add_heading(para, level=2)
                        else:
                            doc.add_paragraph(para)
        
        doc.save(file_path)
    
    def _create_pdf(self, content: str, file_path: str, title: str):
        """Create PDF document from content"""
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            alignment=TA_CENTER,
            fontSize=16,
            spaceAfter=20
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            alignment=TA_JUSTIFY,
            fontSize=10,
            spaceAfter=6
        )
        
        # Build content
        story = []
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Process content sections
        sections = content.split('===')
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
            
            # Check if this is a section header
            if section.isupper() and len(section) < 50:
                story.append(Paragraph(section, heading_style))
            else:
                # Split into paragraphs
                paragraphs = section.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        if para.isupper() and len(para) < 100:
                            story.append(Paragraph(para, heading_style))
                        else:
                            # Escape HTML characters and preserve line breaks
                            para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            para = para.replace('\n', '<br/>')
                            story.append(Paragraph(para, body_style))
        
        doc.build(story)
